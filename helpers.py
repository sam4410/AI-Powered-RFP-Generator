import markdown2
from bs4 import BeautifulSoup
from docx import Document

def convert_md_to_docx(md_text: str, docx_path: str):
    html = markdown2.markdown(md_text)
    soup = BeautifulSoup(html, 'html.parser')
    document = Document()

    for element in soup.descendants:
        if element.name == 'h1':
            document.add_heading(element.text, level=1)
        elif element.name == 'h2':
            document.add_heading(element.text, level=2)
        elif element.name == 'h3':
            document.add_heading(element.text, level=3)
        elif element.name == 'p':
            document.add_paragraph(element.text)
        elif element.name == 'ul':
            for li in element.find_all('li'):
                document.add_paragraph(li.text, style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                document.add_paragraph(li.text, style='List Number')
        elif element.name == 'table':
            rows = element.find_all('tr')
            cols = rows[0].find_all(['th', 'td'])
            table = document.add_table(rows=len(rows), cols=len(cols))
            for i, row in enumerate(rows):
                for j, cell in enumerate(row.find_all(['th', 'td'])):
                    table.cell(i, j).text = cell.text

    document.save(docx_path)
